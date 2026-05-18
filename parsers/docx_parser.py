"""DOCX 解析器 — python-docx 表格提取

DOCX 中的 COD 原始记录通常包含 1-2 个表格:
  - 表格 1: 元数据 + 样品数据 + FAS 标定 + 页脚
  - 表格 2 (可选): 质控结果表

策略: 提取所有表格单元格文本，然后复用 PDF 解析器的 token-stream 逻辑。
"""

import re
from io import BytesIO
from typing import Optional

from models.record import CODRecord
from parsers.field_extractor import safe_float, safe_date, parse_reported_cod, classify_samples


def parse_docx_bytes(file_bytes: bytes, filename: str) -> CODRecord:
    """从 DOCX 字节流解析 COD 记录"""
    from docx import Document
    doc = Document(BytesIO(file_bytes))
    return _extract(doc)


def parse_docx(file_path: str) -> CODRecord:
    """从 DOCX 文件路径解析 COD 记录"""
    from docx import Document
    doc = Document(file_path)
    return _extract(doc)


def _extract(doc) -> CODRecord:
    rec = CODRecord()

    # 提取所有表格为文本行流
    all_lines = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            all_lines.extend(cells)

    # 检测是否有 QC 表 (第二个表格或包含 "质控结果表" 的行)
    qc_start = None
    for i, line in enumerate(all_lines):
        if '质控结果表' in line:
            qc_start = i
            break

    if qc_start is not None:
        main_lines = all_lines[:qc_start]
        qc_lines = all_lines[qc_start:]
        qc_text = '\n'.join(qc_lines)
    else:
        main_lines = all_lines
        qc_text = ""

    main_text = '\n'.join(main_lines)

    # 复用 PDF 解析器的提取逻辑
    from parsers.pdf_parser import (
        _extract_metadata, _extract_instruments, _extract_samples,
        _extract_fas, _extract_footer, _extract_qc, _detect_level,
    )

    _extract_metadata(main_text, rec)
    _extract_instruments(main_text, rec)
    _extract_samples(main_text, rec)
    _extract_fas(main_text, rec)
    _extract_footer(main_text, rec)
    if qc_text:
        _extract_qc(qc_text, rec)

    classify_samples(rec.samples)
    _detect_level(rec)

    return rec
